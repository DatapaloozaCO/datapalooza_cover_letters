# importing required modules
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from datetime import datetime
from bs4 import BeautifulSoup
from requests.auth import HTTPProxyAuth
import requests
import openai 
import constants as ct


def get_completion_from_messages(messages,
                                 model="gpt-3.5-turbo",
                                 temperature=0,
                                 max_tokens=None):
    response = openai.chat.completions.create(
      model=model,
      messages=messages,
      temperature=temperature,
      max_tokens=max_tokens)
    return response.choices[0].message.content


def get_request_with_proxies(url: str):
    auth = HTTPProxyAuth(ct.USERNAME, ct.PASSWORD)
    data = requests.get(
        url=url,
        proxies=ct.PROXIES,
        auth=auth,
        headers={"User-Agent": ct.USER_AGENT},
        timeout=ct.REQUEST_TIMEOUT)
    return data


def get_html_from_url(url: str) -> str:
    """Download HTML file and return a string object"""
    data = get_request_with_proxies(url)
    return data.content


def get_soup_from_url(url: str) -> BeautifulSoup:
    """Download HTML file and return a BeautifulSoup object using
    get_html_from_url function"""
    soup = get_html_from_url(url)
    soup = BeautifulSoup(soup, "html.parser")
    return soup


def generate_cover_letter_text(
        company_url: str,
        curriculum_vitalie: str,
        candidate_instructions: str,
        job_information: str, 
        candidate_name: str) -> str:

    # Get company information
    soup = get_soup_from_url(company_url)
    company_information = soup.get_text(separator=' ', strip=True)

    # Messages that will be pass to ChatGPT
    messages = [
      {'role': 'system', 'content': ct.system_message},
      {'role': 'user', 'content': f"{ct.delimiter} CURRICULUM VITALIE: {curriculum_vitalie} {ct.delimiter}"},
      {'role': 'user', 'content': f"{ct.delimiter} COMPANY INFORMATION {company_information} {ct.delimiter}"},
      {'role': 'user', 'content': f"{ct.delimiter} CANDIDATE INSTRUCTIONS: {candidate_instructions} {ct.delimiter}"},
      {'role': 'user', 'content': f"{ct.delimiter} JOB INFORMATION: {job_information} {ct.delimiter}"},
      {'role': 'user', 'content': f"{ct.delimiter} CANDIDATE NAME: {candidate_name} {ct.delimiter}"}
    ]

    cover_letter_text = get_completion_from_messages(messages)
    return cover_letter_text


def generate_cover_letter_doc(
        your_name: str,
        your_city: str,
        your_country: str,
        your_email: str,
        cover_letter_text: str,
        file_path: str):

    # Create a new Document
    doc = Document()

    # Add a header with subject name
    header = doc.sections[0].header
    paragraph = header.paragraphs[0].add_run(your_name)
    paragraph.font.size = Pt(16)
    paragraph.bold = True

    # Add the date
    date = datetime.now().strftime('%B %d, %Y')

    # Add the header paragraph with your address, email, and date
    header_paragraph = f"{your_country}, {your_city}\n{your_email}\n{date}"
    doc.add_paragraph(header_paragraph)

    # Add document paragraph
    doc.add_paragraph(cover_letter_text)

    # Save the document to a file
    doc.save(file_path)


def pdf_to_text(file_path: str) -> str:
    reader = PdfReader(file_path)

    # printing number of pages in pdf file
    print(len(reader.pages))

    # getting a specific page from the pdf file
    page = reader.pages[0]

    # extracting text from page
    text = page.extract_text()
    return text


def download_file(file_path):
    with open(file_path, "rb") as file:
        file_content = file.read()
    return file_content

import streamlit as st

def reviews_section():
    st.header("User Reviews ⭐")

    reviews = [
        {
            "name": "Emily Johnson",
            "rating": 5,
            "comment": "This cover letter generator is a game-changer! It saved me so much time, and the results were impressive. Highly recommended!",
        },
        {
            "name": "David Chen",
            "rating": 5,
            "comment": "Five stars! I used this service for my job applications, and the cover letters it generated were professional and tailored. It really speeds up the process!",
        },
        {
            "name": "Sarah Rodriguez",
            "rating": 5,
            "comment": "I can't believe how easy it is to create customized cover letters with this tool. It's a must-have for anyone on the job hunt!",
        },
        {
            "name": "Jessica Patel",
            "rating": 5,
            "comment": "This service is a time-saver! I've recommended it to all my friends. It takes the hassle out of writing cover letters and helps you put your best foot forward.",
        },
        {
            "name": "Michael Thompson",
            "rating": 5,
            "comment": "Absolutely fantastic! The cover letters generated are top-notch. I've already landed interviews thanks to the professional touch this service adds.",
        },
    ]

    for review in reviews:
        st.subheader(f"{review['name']} - ⭐⭐⭐⭐⭐")
        st.write(f"*{review['comment']}*\n")